from openai import OpenAI
import os
from docx import Document
import io
import json
import tempfile
import logging
from docx2pdf import convert
from app.utils.mistral_ocr_util import process_pdf_with_mistral_ocr, extract_text_with_pypdf

# Get logger for this module
logger = logging.getLogger("text_processing")

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def extract_text_from_file(file):
    """
    Extract text from different file types with enhanced PDF handling
    
    Args:
        file: File object from request
        
    Returns:
        Extracted text
    """
    filename = file.filename.lower()
    content = file.read()
    
    logger.info(f"Processing file: {filename}")
    
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, os.path.basename(filename))
            
            # Write the file content to a temporary file
            with open(temp_file_path, 'wb') as f:
                f.write(content)
            
            logger.info(f"File saved to temporary location: {temp_file_path}")
            
            try:
                if filename.endswith('.pdf'):
                    logger.info("Detected PDF file, using Mistral OCR")
                    text = process_pdf_with_mistral_ocr(temp_file_path)
                elif filename.endswith('.docx'):
                    logger.info("Detected DOCX file, converting to PDF")
                    try:
                        pdf_path = os.path.join(temp_dir, os.path.splitext(os.path.basename(filename))[0] + '.pdf')
                        convert(temp_file_path, pdf_path)
                        logger.info(f"DOCX converted to PDF: {pdf_path}")
                        text = process_pdf_with_mistral_ocr(pdf_path)
                    except Exception as docx_error:
                        logger.warning(f"DOCX conversion failed: {str(docx_error)}")
                        logger.info("Falling back to direct DOCX extraction")
                        doc = Document(temp_file_path)
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif filename.endswith('.txt'):
                    logger.info("Detected TXT file, reading directly")
                    with open(temp_file_path, 'r', encoding='utf-8', errors='replace') as f:
                        text = f.read()
                else:
                    error_msg = f"Unsupported file type: {filename}"
                    logger.error(error_msg)
                    raise ValueError(error_msg)
            except Exception as extraction_error:
                logger.warning(f"Primary extraction failed: {str(extraction_error)}")
                logger.info("Attempting fallback extraction methods")
                
                if filename.endswith('.pdf'):
                    logger.info("Falling back to PyPDF2 for PDF extraction")
                    text = extract_text_with_pypdf(temp_file_path)
                elif filename.endswith('.docx'):
                    logger.info("Falling back to direct DOCX extraction")
                    doc = Document(temp_file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif filename.endswith('.txt'):
                    logger.info("Falling back to basic text file reading")
                    with open(temp_file_path, 'r', encoding='utf-8', errors='replace') as f:
                        text = f.read()
                else:
                    error_msg = f"Unsupported file type: {filename}"
                    logger.error(error_msg)
                    raise ValueError(error_msg)
            
            logger.info(f"Successfully extracted {len(text)} characters of text")
            return text
    except Exception as e:
        logger.error(f"Error processing {filename}: {str(e)}", exc_info=True)
        raise Exception(f"Error processing {filename}: {str(e)}")

def analyze_document(file):
    """
    Analyze a single document and return structured summary and key facts
    
    Args:
        file: File object from request
        
    Returns:
        Dictionary with filename, summary, and key facts
    """
    logger.info(f"Analyzing document: {file.filename}")
    
    try:
        # Extract text based on file type
        content = extract_text_from_file(file)
        logger.info(f"Text extraction complete. Content length: {len(content)} characters")
        
        # Define the function schema for structured output
        functions = [
            {
                "name": "document_analysis",
                "description": "Analyze a legal document and provide a summary and key facts",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "summary": {
                            "type": "string",
                            "description": "A comprehensive summary of the document"
                        },
                        "key_facts": {
                            "type": "array",
                            "description": "List of key facts or points from the document",
                            "items": {
                                "type": "string"
                            }
                        }
                    },
                    "required": ["summary", "key_facts"]
                }
            }
        ]
        
        # Enhanced system prompt for document analysis
        system_prompt = """You are an expert legal document analyzer with years of experience in legal document review.

Your task is to analyze the provided legal document and extract:
1. A comprehensive yet concise summary that captures the essence of the document
2. Key facts and important points that would be relevant to a legal professional

Guidelines for your analysis:
- Focus on identifying legally significant elements (parties, obligations, dates, conditions)
- Highlight potential legal issues or ambiguities in the document
- Organize information in a logical, structured manner
- Use clear, precise language appropriate for legal professionals
- Prioritize the most important information

Your analysis should be thorough yet accessible, providing valuable insights that would help a legal professional quickly understand the document's significance.
"""
        
        logger.info("Sending content to GPT-4 for analysis")
        # Get analysis using function calling
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content}
            ],
            functions=functions,
            function_call={"name": "document_analysis"}
        )
        
        # Parse the function call result
        function_call = response.choices[0].message.function_call
        analysis_result = json.loads(function_call.arguments)
        
        logger.info("Document analysis complete")
        logger.debug(f"Summary length: {len(analysis_result['summary'])} characters")
        logger.debug(f"Key facts count: {len(analysis_result['key_facts'])}")
        
        return {
            "filename": file.filename,
            "summary": analysis_result["summary"],
            "key_facts": analysis_result["key_facts"]
        }
            
    except Exception as e:
        logger.error(f"Error analyzing {file.filename}: {str(e)}", exc_info=True)
        raise Exception(f"Error analyzing {file.filename}: {str(e)}")

# Keep these functions for backward compatibility
def extract_key_points(content):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Extract and list the key points from this legal document."},
            {"role": "user", "content": content}
        ]
    )
    return response.choices[0].message.content

def identify_missing_clauses(content):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "Identify any standard clauses that are missing from this legal document."},
            {"role": "user", "content": content}
        ]
    )
    return response.choices[0].message.content

def structure_legal_inquiry(question):
    """
    Structure an unclear legal question into a well-formatted inquiry
    
    Args:
        question: The original legal question from the user
        
    Returns:
        Dictionary with structured question and suggested categories
    """
    logger.info("Structuring legal inquiry")
    logger.debug(f"Original question: {question}")
    
    try:
        # Define the function schema for structured output
        functions = [
            {
                "name": "structure_legal_question",
                "description": "Structure an unclear legal question into a well-formatted inquiry",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "structured_question": {
                            "type": "string",
                            "description": "The restructured, clear legal question"
                        },
                        "legal_issues": {
                            "type": "array",
                            "description": "List of legal issues identified in the question",
                            "items": {
                                "type": "string"
                            }
                        },
                        "suggested_categories": {
                            "type": "array",
                            "description": "Categories that this legal question falls under",
                            "items": {
                                "type": "string"
                            }
                        },
                        "missing_information": {
                            "type": "array",
                            "description": "Information that would be helpful to provide a complete answer",
                            "items": {
                                "type": "string"
                            }
                        }
                    },
                    "required": ["structured_question", "legal_issues", "suggested_categories", "missing_information"]
                }
            }
        ]
        
        # Enhanced system prompt for legal inquiry structuring
        system_prompt = """You are an expert legal consultant specializing in transforming unclear legal questions into structured inquiries.

Your task is to:
1. Rewrite the question in a clear, precise, and legally appropriate format
2. Identify the specific legal issues and principles at play
3. Categorize the question into relevant legal domains (e.g., contract law, family law, property law)
4. Note any critical missing information that would be necessary for a complete legal analysis

Guidelines for structuring:
- Use legally precise terminology while remaining accessible
- Separate multiple issues if present in the original question
- Identify the jurisdiction relevance if applicable
- Highlight any assumptions you're making based on the limited information
- Maintain the original intent while adding clarity and structure

Your goal is to transform vague or emotional inquiries into focused legal questions that can be effectively addressed by legal professionals.
"""
        
        logger.info("Sending question to GPT-4 for structuring")
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            functions=functions,
            function_call={"name": "structure_legal_question"}
        )
        
        # Parse the function call result
        function_call = response.choices[0].message.function_call
        result = json.loads(function_call.arguments)
        
        logger.info("Question structuring complete")
        logger.debug(f"Structured result: {result}")
        
        return result
            
    except Exception as e:
        logger.error(f"Error structuring legal inquiry: {str(e)}", exc_info=True)
        raise Exception(f"Error structuring legal inquiry: {str(e)}")

def get_legal_faq_response(question):
    """
    Provide a concise answer to a legal FAQ question
    
    Args:
        question: The legal question from the user
        
    Returns:
        Dictionary with answer and optional references
    """
    logger.info("Processing legal FAQ question")
    logger.debug(f"Question: {question}")
    
    try:
        # Define the function schema for structured output
        functions = [
            {
                "name": "legal_faq_response",
                "description": "Provide a concise answer to a legal question",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "answer": {
                            "type": "string",
                            "description": "A clear, accurate answer to the legal question"
                        },
                        "references": {
                            "type": "array",
                            "description": "References to relevant laws, statutes, or legal principles (optional)",
                            "items": {
                                "type": "string"
                            }
                        }
                    },
                    "required": ["answer"]
                }
            }
        ]
        
        # Enhanced system prompt for legal FAQ responses
        system_prompt = """You are a knowledgeable legal information provider with expertise across multiple areas of law.

Your task is to:
1. Provide a clear, accurate, and concise answer to the legal question
2. Include relevant references to specific laws, statutes, or legal principles when appropriate
3. Acknowledge jurisdictional differences when they significantly impact the answer

Guidelines for your response:
- Balance accuracy with accessibility - use plain language while maintaining legal precision
- Focus on providing factual information rather than advice for specific situations
- Clarify when legal outcomes depend on specific circumstances or jurisdictional variations
- Include a brief disclaimer that your information is for general educational purposes only
- Prioritize clarity and brevity while ensuring the answer is complete

Your goal is to provide helpful, accurate legal information that educates the user while clearly indicating the limits of general legal information.
"""
        
        logger.info("Sending question to GPT-4 for legal FAQ response")
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            functions=functions,
            function_call={"name": "legal_faq_response"}
        )
        
        # Parse the function call result
        function_call = response.choices[0].message.function_call
        result = json.loads(function_call.arguments)
        
        logger.info("Legal FAQ response generated")
        logger.debug(f"Answer length: {len(result['answer'])} characters")
        
        return result
            
    except Exception as e:
        logger.error(f"Error generating legal FAQ response: {str(e)}", exc_info=True)
        raise Exception(f"Error generating legal FAQ response: {str(e)}")
